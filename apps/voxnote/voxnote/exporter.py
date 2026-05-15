"""Export a session to markdown / docx / SRT.

The exporters are intentionally pure functions so they can be re-used from the
CLI, the JS bridge, and tests.
"""
from __future__ import annotations

import datetime as dt
from pathlib import Path

from .storage import Segment, Session, Storage, Summary


def export(storage: Storage, session_id: str, fmt: str, out_dir: Path) -> Path:
    fmt = fmt.lower()
    sess = storage.get_session(session_id)
    if not sess:
        raise ValueError(f"unknown session {session_id!r}")
    segments = storage.list_segments(session_id)
    summary = storage.get_latest_summary(session_id)

    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    safe_title = _safe_filename(sess.title)
    if fmt == "md":
        return _export_md(sess, segments, summary, out_dir / f"{safe_title}.md")
    if fmt == "docx":
        return _export_docx(sess, segments, summary, out_dir / f"{safe_title}.docx")
    if fmt == "srt":
        return _export_srt(segments, out_dir / f"{safe_title}.srt")
    raise ValueError(f"unknown format {fmt!r}; choose md | docx | srt")


# -- markdown ----------------------------------------------------------------

def _export_md(sess: Session, segments: list[Segment], summary: Summary | None, path: Path) -> Path:
    lines: list[str] = []
    lines.append(f"# {sess.title}")
    lines.append("")
    started = dt.datetime.fromtimestamp(sess.started_at).strftime("%Y-%m-%d %H:%M")
    duration = _hms(sess.duration_ms or 0)
    lines.append(f"**Recorded:** {started} ({duration})")
    lines.append("")
    if summary:
        lines.append("## Summary")
        lines.append(summary.summary_md.strip())
        lines.append("")
        if summary.todos:
            lines.append("## Action items")
            for t in summary.todos:
                lines.append(f"- [ ] {t}")
            lines.append("")
    lines.append("## Transcript")
    for seg in segments:
        ts = _hms(seg.start_ms)
        speaker = f" **{seg.speaker}**" if seg.speaker else ""
        lines.append(f"`{ts}`{speaker} {seg.text}")
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


# -- docx --------------------------------------------------------------------

def _export_docx(sess: Session, segments: list[Segment], summary: Summary | None, path: Path) -> Path:
    try:
        from docx import Document  # type: ignore
    except ImportError as err:
        raise RuntimeError(
            "DOCX export needs python-docx. Install voxnote with [export] extra."
        ) from err

    doc = Document()
    doc.add_heading(sess.title, level=1)
    started = dt.datetime.fromtimestamp(sess.started_at).strftime("%Y-%m-%d %H:%M")
    doc.add_paragraph(f"Recorded: {started}  ({_hms(sess.duration_ms or 0)})").italic = True

    if summary:
        doc.add_heading("Summary", level=2)
        for line in summary.summary_md.strip().splitlines():
            doc.add_paragraph(line)
        if summary.todos:
            doc.add_heading("Action items", level=2)
            for t in summary.todos:
                doc.add_paragraph(t, style="List Bullet")

    doc.add_heading("Transcript", level=2)
    for seg in segments:
        para = doc.add_paragraph()
        para.add_run(f"[{_hms(seg.start_ms)}] ").bold = True
        if seg.speaker:
            para.add_run(f"{seg.speaker}: ").italic = True
        para.add_run(seg.text)
    doc.save(str(path))
    return path


# -- srt ---------------------------------------------------------------------

def _export_srt(segments: list[Segment], path: Path) -> Path:
    out: list[str] = []
    for i, seg in enumerate(segments, start=1):
        out.append(str(i))
        out.append(f"{_srt_ts(seg.start_ms)} --> {_srt_ts(seg.end_ms)}")
        out.append(seg.text)
        out.append("")
    path.write_text("\n".join(out), encoding="utf-8")
    return path


# -- helpers -----------------------------------------------------------------

def _hms(ms: int) -> str:
    s, ms_left = divmod(ms, 1000)
    h, s = divmod(s, 3600)
    m, s = divmod(s, 60)
    return f"{h:02d}:{m:02d}:{s:02d}"


def _srt_ts(ms: int) -> str:
    s, ms_left = divmod(ms, 1000)
    h, s = divmod(s, 3600)
    m, s = divmod(s, 60)
    return f"{h:02d}:{m:02d}:{s:02d},{ms_left:03d}"


def _safe_filename(text: str) -> str:
    out = "".join(c if c.isalnum() or c in " -_" else "_" for c in text).strip()
    return (out or "session")[:80]
